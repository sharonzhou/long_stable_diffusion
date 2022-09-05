import argparse
import glob
import os

from docx import Document
from docx.shared import Inches


def dump_images_captions_docx(title, prompts_and_image_paths):
    document = Document()

    document.add_heading(f'{title} images with captions', 0)

    for prompt, image_path in prompts_and_image_paths:
        document.add_paragraph(prompt)
        document.add_picture(image_path, width=Inches(1.25))

    title = title.split('.')[0]
    document.save(f'docx/{title} - illustrations.docx')


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--files", 
        "-f", 
        type=str, 
        required=True, 
        nargs='+', 
        help="File for text"
    )
    args = parser.parse_args()
    os.makedirs('docx/', exist_ok=True)

    files = args.files
    for file in files:
        image_dir_path = f'images/{file}/'
        image_paths = glob.glob(f'{image_dir_path}/*.png')

        prompts = []
        for p in image_paths:
            # get prompt from name of png
            prompt = p.split('/')[-1].split('-')[1]

            # replace underscores with spaces in prompt
            prompt = prompt.replace('_', ' ')
            
            prompts.append(prompt)

        dump_images_captions_docx(file, zip(prompts, image_paths))